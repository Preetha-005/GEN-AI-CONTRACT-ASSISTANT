"""
Document Parser Module
Handles parsing of PDF, DOCX, and TXT files
"""
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import logging

# PDF parsing
import PyPDF2
import pdfplumber

# DOCX parsing
from docx import Document

# Language detection
from langdetect import detect, LangDetectException

import config

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentParser:
    """Parses legal documents and extracts text content"""
    
    def __init__(self):
        self.supported_formats = config.SUPPORTED_FORMATS
        self.max_file_size = config.MAX_FILE_SIZE_MB * 1024 * 1024  # Convert to bytes
    
    def parse_document(self, file_path: str) -> Dict:
        """
        Main method to parse any supported document format
        
        Args:
            file_path: Path to the document
            
        Returns:
            Dictionary containing parsed content and metadata
        """
        path = Path(file_path)
        
        # Validate file
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if path.stat().st_size > self.max_file_size:
            raise ValueError(f"File size exceeds {config.MAX_FILE_SIZE_MB}MB limit")
        
        file_extension = path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Parse based on file type
        if file_extension == '.pdf':
            content = self._parse_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            content = self._parse_docx(file_path)
        elif file_extension == '.txt':
            content = self._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported format: {file_extension}")
        
        # Detect language
        language = self._detect_language(content['text'])
        
        # Extract metadata
        metadata = {
            'filename': path.name,
            'file_type': file_extension,
            'file_size': path.stat().st_size,
            'language': language,
            'page_count': content.get('page_count', 1),
            'word_count': len(content['text'].split()),
            'char_count': len(content['text'])
        }
        
        return {
            'text': content['text'],
            'pages': content.get('pages', [content['text']]),
            'metadata': metadata,
            'tables': content.get('tables', [])
        }
    
    def _parse_pdf(self, file_path: str) -> Dict:
        """Parse PDF file using multiple methods for best results"""
        text_content = []
        tables = []
        
        try:
            # First attempt: Use pdfplumber for better text extraction
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
                    
                    # Extract tables
                    page_tables = page.extract_tables()
                    if page_tables:
                        tables.extend(page_tables)
            
            full_text = "\n\n".join(text_content)
            
            # Fallback to PyPDF2 if pdfplumber fails
            if not full_text.strip():
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text_content = []
                    for page in pdf_reader.pages:
                        text_content.append(page.extract_text())
                    full_text = "\n\n".join(text_content)
            
            return {
                'text': full_text,
                'pages': text_content,
                'page_count': len(text_content),
                'tables': tables
            }
            
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            raise
    
    def _parse_docx(self, file_path: str) -> Dict:
        """Parse DOCX file"""
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            full_text = "\n\n".join(paragraphs)
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                'text': full_text,
                'pages': [full_text],  # DOCX doesn't have pages
                'page_count': 1,
                'tables': tables
            }
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            raise
    
    def _parse_txt(self, file_path: str) -> Dict:
        """Parse plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            return {
                'text': text,
                'pages': [text],
                'page_count': 1,
                'tables': []
            }
            
        except UnicodeDecodeError:
            # Try different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
            
            return {
                'text': text,
                'pages': [text],
                'page_count': 1,
                'tables': []
            }
    
    def _detect_language(self, text: str) -> str:
        """Detect the language of the text"""
        try:
            # Take a sample for language detection (first 1000 chars)
            sample = text[:1000] if len(text) > 1000 else text
            detected_lang = detect(sample)
            
            # Map to supported languages
            if detected_lang in config.SUPPORTED_LANGUAGES:
                return detected_lang
            elif detected_lang in ['hi', 'mr', 'bn']:  # Indian languages
                return 'hi'
            else:
                return 'en'  # Default to English
                
        except LangDetectException:
            logger.warning("Language detection failed, defaulting to English")
            return 'en'
    
    def extract_sections(self, text: str) -> Dict[str, str]:
        """
        Extract common contract sections
        
        Args:
            text: Full contract text
            
        Returns:
            Dictionary of section names and their content
        """
        sections = {}
        
        # Common section patterns in contracts
        section_patterns = [
            r'(?i)(parties|parties to the agreement)',
            r'(?i)(recitals|whereas)',
            r'(?i)(definitions|definitions and interpretation)',
            r'(?i)(term|duration|period of agreement)',
            r'(?i)(scope of work|services|obligations)',
            r'(?i)(payment|compensation|fees)',
            r'(?i)(confidentiality|non-disclosure)',
            r'(?i)(termination|termination provisions)',
            r'(?i)(liability|indemnity|indemnification)',
            r'(?i)(dispute resolution|arbitration)',
            r'(?i)(intellectual property|ip rights)',
            r'(?i)(general provisions|miscellaneous)',
            r'(?i)(governing law|jurisdiction)',
            r'(?i)(warranties|representations)',
        ]
        
        # Split text into potential sections
        lines = text.split('\n')
        current_section = 'Preamble'
        current_content = []
        
        for line in lines:
            # Check if line is a section header
            is_header = False
            for pattern in section_patterns:
                if re.match(pattern, line.strip()):
                    # Save previous section
                    if current_content:
                        sections[current_section] = '\n'.join(current_content)
                    
                    # Start new section
                    current_section = line.strip()
                    current_content = []
                    is_header = True
                    break
            
            if not is_header:
                current_content.append(line)
        
        # Save last section
        if current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
    
    def extract_clauses(self, text: str) -> List[Dict]:
        """
        Extract individual clauses from contract text
        
        Args:
            text: Contract text
            
        Returns:
            List of clause dictionaries
        """
        clauses = []
        
        # Split by numbered clauses (e.g., "1.", "1.1", "Article 1")
        patterns = [
            r'\d+\.\d+\.?\s+[A-Z]',    # 1.1 Title
            r'\d+\.\s+[A-Z]',          # 1. Title
            r'Article\s+\d+',         # Article 1
            r'Clause\s+\d+',          # Clause 1
            r'Section\s+\d+',         # Section 1
        ]
        
        # Combine patterns into a single capturing group
        combined_pattern = f"({'|'.join(patterns)})"
        
        # Split text by clause markers
        parts = re.split(combined_pattern, text)
        
        clause_number = 1
        for i in range(1, len(parts), 2):
            if i + 1 < len(parts):
                header = parts[i]
                content = parts[i + 1].strip() if parts[i + 1] else ""
                
                if len(content) >= config.MIN_CLAUSE_LENGTH:
                    clauses.append({
                        'clause_id': f"C{clause_number:03d}",
                        'clause_number': header.strip(),
                        'content': content,
                        'length': len(content),
                        'word_count': len(content.split())
                    })
                    clause_number += 1
        
        # If no structured clauses found, split by paragraphs
        if not clauses:
            paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
            for idx, para in enumerate(paragraphs, 1):
                if len(para) >= config.MIN_CLAUSE_LENGTH:
                    clauses.append({
                        'clause_id': f"C{idx:03d}",
                        'clause_number': f"Para {idx}",
                        'content': para,
                        'length': len(para),
                        'word_count': len(para.split())
                    })
        
        return clauses


if __name__ == "__main__":
    # Test the parser
    parser = DocumentParser()
    print("Document Parser initialized successfully")
    print(f"Supported formats: {parser.supported_formats}")
